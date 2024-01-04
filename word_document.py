import streamlit as st
from docx import Document

def docx_converter(doc):
    if doc is not None:
        document = Document(doc)

        # Extract text from all paragraphs
        all_text = "\n".join([paragraph.text for paragraph in document.paragraphs])

        # st.write("Number of pages:", len(document.paragraphs))
        return all_text
